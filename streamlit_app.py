#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Dieses Script implementiert eine Agenten-basierte Diskussionsplattform,
bei der verschiedene Agenten interaktiv über Themen kommunizieren.
Dabei wird auch die Google Generative AI API genutzt, um Antworten zu generieren.

Wichtiger Hinweis:
Der Fehler "module 'google.genai' has no attribute 'configure'" entsteht,
weil das Modul google.genai nicht die Methode configure enthält.
Der korrekte Import lautet daher:
    import google.generativeai as genai
"""

import re
import logging
import datetime
import json
import hashlib
import os
import time
import random
from collections import defaultdict
from typing import List, Dict, Tuple, Any, Union
import sqlite3
from jsonschema import validate, ValidationError
from docx import Document
from docx.shared import Inches
import streamlit as st
import tornado

# -----------------------------------------------------------------------------
# Wichtige Änderung: Korrekte Importanweisung für Google Generative AI
# Anstatt "from google import genai" wird das korrekte Modul importiert:
import google.generativeai as genai
# Damit wird auch die Methode "configure" zur Verfügung gestellt.
from google.generativeai.types.generation_types import StopCandidateException

# -----------------------------------------------------------------------------
# Konstanten und Logging-Konfiguration

MODEL_NAME = "gemini-2.0-flash-thinking-exp-01-21"
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

API_SLEEP_SECONDS = 60
API_MAX_RETRIES = 3
SUMMARY_SLEEP_SECONDS = 10
AUDIT_LOG_FILE = "audit_log.txt"
EXPIRATION_TIME_SECONDS = 300
ROLE_PERMISSIONS = {
    "user": ["REQ", "DATA"],
    "admin": ["REQ", "DATA", "CALC", "IF", "AI"]
}
PRIORITY_MAP = {"HIGH": 1, "MEDIUM": 2, "LOW": 3}

USER_DATA_FILE = "user_data.json"
DISCUSSION_DB_FILE = "discussion_data.db"
RATING_DATA_FILE = "rating_data.json"
AGENT_CONFIG_FILE = "agent_config.json"

# -----------------------------------------------------------------------------
# JSON-Schema-Definitionen zur Validierung der Eingabedaten

USER_DATA_SCHEMA = {
    "type": "object",
    "patternProperties": {
        "^[a-zA-Z0-9_-]+$": {
            "type": "object",
            "properties": {
                "password": {"type": "string"}
            },
            "required": ["password"]
        }
    }
}

AGENT_CONFIG_SCHEMA = {
    "type": "array",
    "items": {
        "type": "object",
        "properties": {
            "name": {"type": "string"},
            "personality": {"type": "string", "enum": ["kritisch", "visionär", "konservativ", "neutral", "kreativ", "analytisch", "humorvoll"]},
            "description": {"type": "string"},
            "instruction": {"type": "string"}
        },
        "required": ["name", "personality", "description"]
    }
}

# -----------------------------------------------------------------------------
# Funktionen zum Laden und Speichern von JSON-Daten

def load_json_data(filename: str, schema: dict = None) -> Dict[str, Any]:
    """
    Lädt JSON-Daten aus einer Datei und validiert sie optional gegen ein Schema.
    """
    try:
        with open(filename, "r", encoding="utf-8") as f:
            data = json.load(f)
            if schema:
                validate(instance=data, schema=schema)
            return data
    except FileNotFoundError:
        logging.warning(f"Datei '{filename}' nicht gefunden. Starte mit leeren Daten.")
        return {}
    except json.JSONDecodeError as e:
        logging.error(f"Fehler beim Lesen von '{filename}': Ungültiges JSON-Format. Details: {e}")
        return {}
    except ValidationError as e:
        logging.error(f"Datei '{filename}' entspricht nicht dem erwarteten Schema: {e}")
        return {}

def save_json_data(data: Dict[str, Any], filename: str) -> None:
    """
    Speichert ein Dictionary als formatiertes JSON in eine Datei.
    """
    try:
        with open(filename, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=4)
    except IOError as e:
        logging.error(f"Fehler beim Schreiben in Datei '{filename}': {e}")

# -----------------------------------------------------------------------------
# Funktionen zum Laden und Speichern von Benutzerdaten

def load_user_data() -> Dict[str, Any]:
    return load_json_data(USER_DATA_FILE, USER_DATA_SCHEMA)

def save_user_data(user_data: Dict[str, Any]) -> None:
    save_json_data(user_data, USER_DATA_FILE)

# -----------------------------------------------------------------------------
# Funktionen zum Laden und Speichern von Bewertungsdaten

def load_rating_data() -> Dict[str, Any]:
    return load_json_data(RATING_DATA_FILE)

def save_rating_data(rating_data: Dict[str, Any]) -> None:
    save_json_data(rating_data, RATING_DATA_FILE)

# -----------------------------------------------------------------------------
# Funktionen zur Konfiguration der Agenten

def load_agent_config() -> List[Dict[str, str]]:
    config = load_json_data(AGENT_CONFIG_FILE, AGENT_CONFIG_SCHEMA)
    if not isinstance(config, list):
        logging.error(f"Agentenkonfiguration in '{AGENT_CONFIG_FILE}' ist ungültig oder leer.")
        return []
    return config

# -----------------------------------------------------------------------------
# Hilfsfunktionen zur Passwortverarbeitung

def hash_password(password: str) -> str:
    """
    Erzeugt einen SHA256-Hash des Passworts.
    """
    return hashlib.sha256(password.encode()).hexdigest()

def verify_password(password: str, hashed_password: str) -> bool:
    """
    Vergleicht ein Klartext-Passwort mit einem gehashten Passwort.
    """
    return hash_password(password) == hashed_password

# -----------------------------------------------------------------------------
# Benutzerregistrierung und Login

def register_user(username: str, password: str) -> str:
    """
    Registriert einen neuen Benutzer, sofern der Nutzername und das Passwort den Anforderungen entsprechen.
    """
    if not re.match(r"^[a-zA-Z0-9_-]+$", username):
        return "Ungültiger Nutzername. Nur Buchstaben, Zahlen, '-', '_' erlaubt."
    if len(password) < 8:
        return "Passwort muss mindestens 8 Zeichen lang sein."
    user_data = load_user_data()
    if username in user_data:
        return "Nutzername bereits vergeben."
    user_data[username] = {"password": hash_password(password)}
    save_user_data(user_data)
    return "Registrierung erfolgreich."

def login_user(username: str, password: str) -> Tuple[str, Union[str, None]]:
    """
    Überprüft die Login-Daten und gibt den Login-Status zurück.
    """
    user_data = load_user_data()
    if username in user_data and verify_password(password, user_data[username]["password"]):
        return "Login erfolgreich.", username
    return "Login fehlgeschlagen.", None

# -----------------------------------------------------------------------------
# Datenbankfunktionen für Diskussionen

def create_discussion_table():
    """
    Erstellt die SQLite-Tabelle für Diskussionen, sofern sie nicht existiert.
    """
    conn = sqlite3.connect(DISCUSSION_DB_FILE)
    cursor = conn.cursor()
    cursor.execute("""
        CREATE TABLE IF NOT EXISTS discussions (
            discussion_id TEXT PRIMARY KEY,
            topic TEXT,
            agents TEXT,
            chat_history TEXT,
            summary TEXT,
            user TEXT,
            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    """)
    conn.commit()
    conn.close()

create_discussion_table()

def save_discussion_data_db(discussion_id: str, topic: str, agents: List[str], chat_history: List[Dict], summary: str, user: str = None) -> None:
    """
    Speichert eine Diskussion in der SQLite-Datenbank.
    """
    conn = sqlite3.connect(DISCUSSION_DB_FILE)
    cursor = conn.cursor()
    try:
        cursor.execute("""
            INSERT INTO discussions (discussion_id, topic, agents, chat_history, summary, user)
            VALUES (?, ?, ?, ?, ?, ?)
        """, (discussion_id, topic, json.dumps(agents), json.dumps(chat_history), summary, user))
        conn.commit()
    except sqlite3.Error as e:
        logging.error(f"Datenbankfehler beim Speichern der Diskussion '{discussion_id}': {e}")
        conn.rollback()
    finally:
        conn.close()

def load_discussion_data_db(user: str = None) -> Dict[str, Any]:
    """
    Lädt Diskussionen aus der SQLite-Datenbank; falls ein Benutzer angegeben ist, werden nur dessen Diskussionen geladen.
    """
    conn = sqlite3.connect(DISCUSSION_DB_FILE)
    cursor = conn.cursor()
    discussions = {}
    try:
        if user:
            cursor.execute("SELECT * FROM discussions WHERE user = ?", (user,))
        else:
            cursor.execute("SELECT * FROM discussions")
        rows = cursor.fetchall()
        for row in rows:
            disc_id, topic, agents_json, chat_history_json, summary, user_name, timestamp = row
            agents = json.loads(agents_json) if agents_json else []
            chat_history = json.loads(chat_history_json) if chat_history_json else []
            discussions[disc_id] = {
                "topic": topic,
                "agents": agents,
                "chat_history": chat_history,
                "summary": summary,
                "user": user_name,
                "timestamp": timestamp
            }
    except sqlite3.Error as e:
        logging.error(f"Datenbankfehler beim Laden der Diskussionen: {e}")
    finally:
        conn.close()
    return discussions

# -----------------------------------------------------------------------------
# Funktion zur Bewertung der Agentenantworten

def evaluate_response(response: str) -> str:
    """
    Bewertet die Antwort eines Agenten anhand bestimmter Schlüsselwörter.
    """
    resp_l = response.lower()
    if "wiederhole mich" in resp_l:
        return "schlechte antwort"
    elif "neue perspektive" in resp_l:
        return "gute antwort"
    else:
        return "neutral"

discussion_ratings = defaultdict(lambda: defaultdict(dict), load_rating_data())

def rate_agent_response(discussion_id: str, iteration: int, agent_name: str, rating_type: str) -> None:
    """
    Erhöht den Upvote- oder Downvote-Zähler für einen Agenten in einer bestimmten Diskussionsrunde.
    """
    global discussion_ratings
    if agent_name not in discussion_ratings[discussion_id][iteration]:
        discussion_ratings[discussion_id][iteration][agent_name] = {"upvotes": 0, "downvotes": 0}
    if rating_type == "upvote":
        discussion_ratings[discussion_id][iteration][agent_name]["upvotes"] += 1
    elif rating_type == "downvote":
        discussion_ratings[discussion_id][iteration][agent_name]["downvotes"] += 1
    save_rating_data(discussion_ratings)

# -----------------------------------------------------------------------------
# Funktionen zur Zusammenfassung von PDF-Inhalten und zum API-Aufruf

def generate_pdf_summary_from_bytes(file_bytes: bytes, api_key: str) -> str:
    """
    Erzeugt eine Zusammenfassung des Inhalts einer PDF-Datei mithilfe der Gemini API.
    """
    try:
        mime_type = "application/pdf"
        prompt = "Fasse den Inhalt der PDF zusammen. Achte dabei darauf dass wichtige Daten nicht verloren gehen!"
        # Erzeugt den Inhalt als Liste: Zunächst der Prompt, dann der PDF-Content als Teil
        contents = [prompt, genai.types.Part.from_bytes(data=file_bytes, mime_type=mime_type)]
        response = call_gemini_api(contents, api_key)
        return response.get("response", "Start der Konversation.")
    except Exception as e:
        logging.error("Fehler beim Generieren der PDF-Zusammenfassung:", exc_info=e)
        return "Start der Konversation."

def call_gemini_api(contents: list, api_key: str) -> Dict[str, str]:
    """
    Ruft die Gemini API auf, um Inhalte zu generieren. Hierbei wird eine Konfiguration mit dem API-Schlüssel vorgenommen.
    """
    # Konfiguration der API: Hier wird der API-Schlüssel gesetzt
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(MODEL_NAME)
    retries = 0
    wait_time = 1
    max_wait_time = 60
    while retries < API_MAX_RETRIES:
        try:
            logging.info(f"Sende Anfrage an Gemini: {str(contents)[:100]}... (Versuch {retries + 1})")
            response = model.generate_content(contents=contents)
            if response.text is None or not response.text.strip():
                msg = "Leere Antwort von Gemini API."
                logging.warning(msg)
                return {"response": msg}
            return {"response": response.text}
        except Exception as e:
            err_s = str(e)
            logging.error(f"Gemini API Fehler: {err_s}")
            if "429" in err_s or "Too Many Requests" in err_s:
                retries += 1
                if retries >= API_MAX_RETRIES:
                    return {"response": f"Fehler: Maximale Anzahl an Versuchen erreicht. API-Kontingent wahrscheinlich erschöpft."}
                wait_time = min(wait_time * 2, max_wait_time)
                jitter = random.uniform(0.5, 1.5)
                actual_wait_time = wait_time * jitter
                logging.warning(f"API-Kontingent erschöpft. Warte {actual_wait_time:.2f} Sekunden...")
                time.sleep(actual_wait_time)
            else:
                return {"response": f"Fehler bei Gemini API Aufruf: {err_s}"}
    return {"response": f"Fehler: Maximale Anzahl an Versuchen erreicht ({API_MAX_RETRIES})."}

def generate_summary(text: str, api_key: str) -> str:
    """
    Generiert eine prägnante Zusammenfassung eines gegebenen Texts mithilfe der Gemini API.
    """
    prompt = f"Fasse den folgenden Text prägnant zusammen:\n\n{text}"
    result = call_gemini_api([prompt], api_key)
    if SUMMARY_SLEEP_SECONDS > 0:
        time.sleep(SUMMARY_SLEEP_SECONDS)
    return result.get("response", "Fehler: Keine Zusammenfassung generiert.")

# -----------------------------------------------------------------------------
# Hauptfunktion für die gemeinsame Konversation der Agenten

def joint_conversation_with_selected_agents(
    conversation_topic: str,
    selected_agents: List[Dict[str, str]],
    iterations: int,
    expertise_level: str,
    language: str,
    chat_history: List[Dict[str, str]],
    user_state: str,
    discussion_id: str = None,
    api_key: str = None,
    pdf_file = None
):
    """
    Startet eine interaktive Konversation zwischen ausgewählten Agenten.
    Hierbei wird unter anderem eine API-Anfrage an Gemini gesendet.
    """
    if discussion_id is None:
        discussion_id = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    chat_history_filename = f"chat_history_{discussion_id}.txt"
    active_agents_names = [sa["name"] for sa in selected_agents]
    num_agents = len(active_agents_names)
    agent_outputs = [""] * num_agents
    topic_changed = False
    logging.info(f"Konversation gestartet: {active_agents_names}, iters={iterations}, level={expertise_level}, lang={language}, Diskussions-ID: {discussion_id}, PDF: {pdf_file is not None}")

    # Falls eine PDF bereitgestellt wurde, wird deren Inhalt zusammengefasst
    if pdf_file is not None:
        try:
            pdf_file_bytes = pdf_file.read()
            pdf_file.seek(0)
            initial_summary = generate_pdf_summary_from_bytes(pdf_file_bytes, api_key)
        except Exception as e:
            logging.error(f"Fehler beim Lesen/Zusammenfassen der PDF: {e}", exc_info=e)
            initial_summary = "Start der Konversation."
    else:
        initial_summary = "Start der Konversation."
    current_summary = initial_summary

    for i in range(iterations):
        agent_idx = i % num_agents
        current_agent_name = active_agents_names[agent_idx]
        current_agent_config = next((a for a in selected_agents if a["name"] == current_agent_name), None)
        current_personality = current_agent_config.get("personality", "neutral")
        current_instruction = current_agent_config.get("instruction", "")

        prompt_text = (
            f"Wir führen eine Konversation über: '{conversation_topic}'.\n"
            + ("Zusätzliche Informationen sind in der angehängten PDF-Datei verfügbar.\n" if pdf_file is not None else "")
            + f"Hier ist die Zusammenfassung der bisherigen Diskussion:\n{current_summary}\n\n"
            + f"Iteration {i+1}: Agent {current_agent_name}, bitte antworte. {current_instruction}\n"
        )
        if i > 0:
            prompt_text += f"Der vorherige Agent sagte: {agent_outputs[(agent_idx - 1) % num_agents]}\n"

        # Anpassung des Prompts je nach Persönlichkeit des Agenten
        if current_personality == "kritisch":
            prompt_text += "\nSei kritisch und hinterfrage alles."
        elif current_personality == "visionär":
            prompt_text += "\nSei visionär und denke an die Zukunft."
        elif current_personality == "konservativ":
            prompt_text += "\nSei konservativ und berücksichtige traditionelle Ansätze."
        elif current_personality == "neutral":
            prompt_text += "\nAntworte neutral und objektiv."
        elif current_personality == "kreativ":
            prompt_text += "\nSei kreativ und innovativ in deiner Antwort."
        elif current_personality == "analytisch":
            prompt_text += "\nAnalysiere die Situation und antworte analytisch."
        elif current_personality == "humorvoll":
            prompt_text += "\nLockere die Konversation mit etwas Humor auf."

        prompt_text += f"\n\nAntworte auf {language}."

        # Inhalt für die API-Anfrage – falls eine PDF vorhanden ist, wird diese als Teil gesendet
        contents = [prompt_text]
        if pdf_file is not None:
            contents = [genai.types.Part.from_bytes(data=pdf_file_bytes, mime_type="application/pdf"), f"{prompt_text}"]

        api_resp = call_gemini_api(contents, api_key)
        agent_output = api_resp.get("response", f"Keine Antwort von {current_agent_name}")
        agent_outputs[agent_idx] = agent_output

        chat_history.append({
            "role": "user",
            "content": f"Agent {current_agent_name} (Iteration {i + 1}): Thema {conversation_topic}, Zusammenfassung bis Runde {i}: {current_summary}, PDF: {'vorhanden' if pdf_file is not None else 'nicht vorhanden'}"
        })
        chat_history.append({
            "role": "assistant",
            "content": f"Antwort von Agent {current_agent_name} (Iteration {i+1}):\n{agent_output}"
        })

        try:
            with open(chat_history_filename, "a", encoding="utf-8") as f:  # Appending an existing chat log
                f.write(f"Iteration {i+1}, Agent: {current_agent_name}, Persönlichkeit: {current_personality}\n")
                f.write(f"Prompt: {prompt_text}\n")
                f.write(f"Antwort: {agent_output}\n")
                f.write("-" * 50 + "\n")
        except IOError as e:
            logging.error(f"Fehler beim Schreiben in Chatverlauf-Datei '{chat_history_filename}': {e}")

        new_summary_input = f"Bisherige Zusammenfassung:\n{current_summary}\n\nNeue Antwort von {current_agent_name}:\n{agent_output}"
        current_summary = generate_summary(new_summary_input, api_key)
        time.sleep(API_SLEEP_SECONDS)

        qual = evaluate_response(agent_output)
        if qual == "schlechte antwort":
            logging.info(f"{current_agent_name} => 'schlechte antwort', retry ...")
            retry_contents = ["Versuche eine kreativere Antwort."]
            if pdf_file is not None:
                retry_contents = [genai.types.Part.from_bytes(data=pdf_file_bytes, mime_type="application/pdf"), "Versuche eine kreativere Antwort."]
            retry_resp = call_gemini_api(retry_contents, api_key)
            retry_output = retry_resp.get("response", f"Keine Retry-Antwort von {current_agent_name}")
            if "Fehler bei Gemini API Aufruf" not in retry_output:
                agent_output = retry_output
            agent_outputs[agent_idx] = agent_output

        logging.info(f"Antwort Agent {current_agent_name} (i={i+1}): {agent_output[:50]}...")
        formatted_output_chunk = (
            f"**Iteration {i+1}: Agent {current_agent_name} ({current_personality})**\n\n"
            f"{agent_output}\n\n"
            "---\n\n"
        )
        yield chat_history, formatted_output_chunk, discussion_id, (i + 1), current_agent_name

        if i > iterations * 0.6 and agent_output == agent_outputs[(agent_idx - 1) % num_agents] and not topic_changed:
            new_topic = "Neues Thema: KI-Trends 2026"
            contents = [new_topic]
            if pdf_file is not None:
                contents = [genai.types.Part.from_bytes(data=pdf_file_bytes, mime_type="application/pdf"), f"{new_topic}"]
            agent_outputs = [new_topic] * num_agents
            topic_changed = True

    final_summary_input = "Gesamter Chatverlauf:\n" + "\n".join(
        [f"{m['role']}: {m['content']}" for m in chat_history]
    )
    final_summary = generate_summary(final_summary_input, api_key)
    chat_history.append({
        "role": "assistant",
        "content": f"**Gesamtzusammenfassung**:\n{final_summary}"
    })

    if user_state:
        save_discussion_data_db(discussion_id, conversation_topic, active_agents_names, chat_history, final_summary, user_state)
        logging.info(f"Diskussion {discussion_id} für {user_state} in Datenbank gespeichert.")
    else:
        logging.info("Keine Speicherung in Datenbank, kein Benutzer eingeloggt.")

    final_text = agent_outputs[-1]
    chat_history.append({
        "role": "assistant",
        "content": f"Finale Aussage:\n{final_text}"
    })
    logging.info(f"Finale Aussage: {final_text}")
    yield chat_history, final_summary, discussion_id, None, None

# -----------------------------------------------------------------------------
# Funktion zum Speichern des Chatverlaufs als Word-Dokument

def save_chat_as_word(chat_history: List[Dict], discussion_id: str) -> Union[str, None]:
    """
    Speichert den Chatverlauf in einem Word-Dokument.
    """
    document = Document()
    document.add_heading(f'CipherCore Agenten-Diskussion: {discussion_id}', level=1)
    for message in chat_history:
        role = message['role']
        content = message['content']
        if role == 'user':
            document.add_paragraph("Nutzer:", style='List Bullet').add_run(f" {content}").bold = True
        elif role == 'assistant':
            agent_name_match = re.search(r'Agent (.*?)\s', content)
            agent_name = agent_name_match.group(1) if agent_name_match else "Agent"
            p = document.add_paragraph(f"{agent_name}:", style='List Bullet')
            p.add_run(f" {content.split(':\n', 1)[1] if ':\n' in content else content}")
    filename = f"CipherCore_Diskussion_{discussion_id}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    try:
        document.save(filename)
        logging.info(f"Word-Datei '{filename}' erfolgreich gespeichert.")
        return filename
    except Exception as e:
        logging.error(f"Fehler beim Speichern der Word-Datei: {e}")
        return None

# -----------------------------------------------------------------------------
# Hauptfunktion (Streamlit-Webapp)

def main():
    """
    Startet die Streamlit-Webapplikation, welche die Agenten-Diskussion steuert.
    """
    st.set_page_config(page_title="CipherCore Agenten-Konversation", page_icon="🤖")
    st.title("CipherCore Agenten-Konversation")
    st.markdown("Willkommen bei CipherCore! Ihre Plattform für innovative KI-gestützte Diskussionen.")

    # Initialisierung der Session-States
    if 'user_state' not in st.session_state:
        st.session_state['user_state'] = None
    if 'chat_history' not in st.session_state:
        st.session_state['chat_history'] = []
    if 'discussion_id' not in st.session_state:
        st.session_state['discussion_id'] = None
    if 'rating_info' not in st.session_state:
        st.session_state['rating_info'] = {}
    if 'formatted_output_text' not in st.session_state:
        st.session_state['formatted_output_text'] = ""
    if 'api_key' not in st.session_state:
        st.session_state['api_key'] = None
    if 'pdf_file' not in st.session_state:
        st.session_state['pdf_file'] = None

    st.sidebar.header("API-Schlüssel & Benutzer")
    api_key = st.sidebar.text_input("Gemini API-Schlüssel:", type="password")
    if not api_key and not st.session_state['api_key']:
        st.sidebar.warning("Bitte API-Schlüssel eingeben.")
    elif api_key:
        st.session_state['api_key'] = api_key

    with st.sidebar.expander("Login / Registrierung", expanded=False):
        col1, col2 = st.columns(2)
        with col1:
            st.subheader("Login")
            username_login = st.text_input("Nutzername", key="username_login_sidebar")
            password_login = st.text_input("Passwort", type="password", key="password_login_sidebar")
            login_btn = st.button("Login", key="login_btn_sidebar")
            login_status = st.empty()
            if login_btn:
                msg, logged_in_user = login_user(username_login, password_login)
                if logged_in_user:
                    st.session_state['user_state'] = logged_in_user
                    login_status.success(msg)
                    st.rerun()
                else:
                    login_status.error(msg)
        with col2:
            st.subheader("Registrierung")
            username_register = st.text_input("Nutzername", key="username_register_sidebar")
            password_register = st.text_input("Passwort", type="password", key="password_register_sidebar")
            register_btn = st.button("Registrieren", key="register_btn_sidebar")
            register_status = st.empty()
            if register_btn:
                msg = register_user(username_register, password_register)
                register_status.info(msg)
    if st.session_state['user_state']:
        st.sidebar.success(f"Eingeloggt als: {st.session_state['user_state']}")
        if st.sidebar.button("Logout"):
            st.session_state['user_state'] = None
            st.session_state['chat_history'] = []
            st.session_state['discussion_id'] = None
            st.session_state['formatted_output_text'] = ""
            st.rerun()

    st.markdown("---")

    agent_config_data = load_agent_config()
    agent_selections = {}
    st.subheader("Agenten Auswahl")
    with st.expander("Agenten Konfiguration anzeigen/verstecken", expanded=False):
        agent_cols = st.columns(3)
        for idx, agent_data in enumerate(agent_config_data):
            with agent_cols[idx % 3]:
                agent_selections[agent_data["name"]] = {
                    "selected": st.checkbox(agent_data["name"], value=False),
                    "personality": st.selectbox(
                        f"Persönlichkeit für {agent_data['name']}",
                        ["kritisch", "visionär", "konservativ", "neutral", "kreativ", "analytisch", "humorvoll"],
                        index=["kritisch", "visionär", "konservativ", "neutral", "kreativ", "analytisch", "humorvoll"].index(agent_data["personality"]),
                        key=f"personality_{agent_data['name']}"
                    ),
                    "instruction": st.text_area(f"Spezielle Instruktion für {agent_data['name']}", value=agent_data.get("instruction", ""), key=f"instruction_{agent_data['name']}")
                }

    topic_input = st.text_input("Diskussionsthema", value="Wie verbessern wir die Kundenbindung durch KI?")
    iteration_slider = st.slider("Anzahl Gesprächsrunden", 1, 50, value=15, step=1)
    level_radio = st.radio("Experten-Level", ["Beginner", "Fortgeschritten", "Experte"], horizontal=True)
    lang_radio = st.radio("Sprache", ["Deutsch", "Englisch", "Französisch", "Spanisch"], horizontal=True, index=0)

    st.subheader("PDF-Datei (optional)")
    uploaded_file = st.file_uploader("PDF hochladen für Kontext", type="pdf")
    if uploaded_file is not None:
        st.session_state['pdf_file'] = uploaded_file
        st.success("PDF hochgeladen!")

    with st.expander("Gespeicherte Diskussionen (Login erforderlich)", expanded=False):
        if st.session_state['user_state']:
            load_disc_btn = st.button("Gespeicherte Diskussionen laden")
            saved_discussions = st.empty()
            if load_disc_btn:
                disc_data = load_discussion_data_db(st.session_state['user_state'])
                if disc_data:
                    st.write("Wähle eine Diskussion zum Anzeigen:")
                    selected_discussion_id = st.selectbox("Diskussions-ID", options=list(disc_data.keys()))
                    if selected_discussion_id:
                        discussion = disc_data[selected_discussion_id]
                        st.session_state['chat_history'] = discussion['chat_history']
                        st.session_state['discussion_id'] = selected_discussion_id
                        st.session_state['formatted_output_text'] = ""  # Clear formatted output when loading
                        st.rerun()  # Rerun to display loaded chat
                else:
                    saved_discussions.info("Keine gespeicherten Diskussionen gefunden.")
        else:
            st.info("Bitte logge dich ein, um gespeicherte Diskussionen zu sehen.")

    start_btn = st.button("Konversation starten", type="primary")

    st.subheader("Agenten-Konversation")
    chat_display_area = st.container()  # Use a container for chat messages

    st.subheader("Formatierter Output")
    formatted_output_area = st.empty()

    rating_col1, rating_col2, rating_col3 = st.columns([1, 1, 3])
    rating_label = rating_col3.empty()

    if st.session_state['rating_info'].get("iteration") is not None:
        with rating_col1:
            upvote_btn = st.button("👍 Upvote")
        with rating_col2:
            downvote_btn = st.button("👎 Downvote")

        if upvote_btn:
            did = st.session_state['rating_info'].get("discussion_id")
            itn = st.session_state['rating_info'].get("iteration")
            agn = st.session_state['rating_info'].get("agent_name")
            if did and itn and agn:
                rate_agent_response(did, itn, agn, "upvote")
                rating_label.success("👍 Upvote gegeben")
            else:
                rating_label.error("Fehler beim Upvote (fehlende Daten).")
        if downvote_btn:
            did = st.session_state['rating_info'].get("discussion_id")
            itn = st.session_state['rating_info'].get("iteration")
            agn = st.session_state['rating_info'].get("agent_name")
            if did and itn and agn:
                rate_agent_response(did, itn, agn, "downvote")
                rating_label.success("👎 Downvote gegeben")
            else:
                rating_label.error("Fehler beim Downvote (fehlende Daten).")

    save_col1, save_col2 = st.columns(2)
    with save_col1:
        save_btn = st.button("Diskussion speichern (Login)")
        save_status = st.empty()
        if save_btn:
            if st.session_state['user_state']:
                active_agents_names = [agent['name'] for agent in agent_config_data if agent_selections[agent['name']]['selected']]
                save_discussion_data_db(st.session_state['discussion_id'], topic_input, active_agents_names, st.session_state['chat_history'], "Manuell gespeichert", st.session_state['user_state'])
                save_status.success("Diskussion in Datenbank gespeichert.")
            else:
                save_status.warning("Bitte zuerst einloggen.")
    with save_col2:
        word_save_btn = st.button("Chat als Word speichern")
        if word_save_btn:
            if st.session_state['discussion_id']:
                word_filename = save_chat_as_word(st.session_state['chat_history'], st.session_state['discussion_id'])
                if word_filename:
                    with open(word_filename, "rb") as file:
                        st.download_button(
                            label="Word-Datei herunterladen",
                            data=file,
                            file_name=word_filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                else:
                    st.error("Fehler beim Erstellen der Word-Datei.")
            else:
                st.warning("Diskussions-ID fehlt. Starte zuerst eine Konversation.")

    if start_btn and st.session_state['api_key']:
        selected_agents = [
            {"name": agent, "personality": agent_selections[agent]["personality"], "instruction": agent_selections[agent]["instruction"]}
            for agent in agent_selections if agent_selections[agent]["selected"]
        ]
        if not selected_agents:
            st.warning("Bitte wähle mindestens einen Agenten aus.")
        else:
            st.session_state['chat_history'] = []
            st.session_state['formatted_output_text'] = ""
            st.session_state['discussion_id'] = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            st.session_state['rating_info'] = {}

            chat_display_area.empty()  # Clear previous chat messages
            formatted_output_area.empty()  # Clear formatted output

            try:
                with st.spinner("Konversation wird gestartet..."):
                    agent_convo = joint_conversation_with_selected_agents(
                        conversation_topic=topic_input,
                        selected_agents=selected_agents,
                        iterations=iteration_slider,
                        expertise_level=level_radio,
                        language=lang_radio,
                        chat_history=[],
                        user_state=st.session_state['user_state'],
                        discussion_id=st.session_state['discussion_id'],
                        api_key=st.session_state['api_key'],
                        pdf_file=st.session_state.get('pdf_file')
                    )
                    for updated_hist, chunk_text, disc_id, iteration_num, agent_n in agent_convo:
                        st.session_state['discussion_id'] = disc_id
                        st.session_state['rating_info']["discussion_id"] = disc_id
                        st.session_state['rating_info']["iteration"] = iteration_num
                        st.session_state['rating_info']["agent_name"] = agent_n

                        if updated_hist and len(updated_hist) > len(st.session_state['chat_history']):
                            new_messages = updated_hist[len(st.session_state['chat_history']):]
                            for message in new_messages:
                                st.session_state['chat_history'].append(message)
                                with chat_display_area:  # Display messages in the container
                                    with st.chat_message(message["role"]):
                                        st.markdown(message["content"])
                            st.session_state['formatted_output_text'] += chunk_text
                            formatted_output_area.markdown(st.session_state['formatted_output_text'])  # Update formatted output area

            except (tornado.websocket.WebSocketClosedError, tornado.iostream.StreamClosedError) as e:
                st.error(f"Verbindungsfehler: {e}. Bitte später erneut versuchen.")
            except StopCandidateException as e:
                st.error(f"Konversation wurde unerwartet beendet: {e}")
            except Exception as e:
                st.error(f"Unerwarteter Fehler: {e}")
    elif start_btn and not st.session_state['api_key']:
        st.error("Bitte gib einen API-Schlüssel ein, um die Konversation zu starten.")


if __name__ == "__main__":
    main()
